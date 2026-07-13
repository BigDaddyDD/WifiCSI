"""Build a plain progress-update Word doc summarizing recent CSI findings.

No theme, no styling flourishes: Times New Roman 12 pt, black text, embeds
figures 1-6 from docs/model_report_assets. Run: python slides_assets/build_progress_doc.py
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ASSETS = os.path.join(ROOT, "docs", "model_report_assets")
OUT = os.path.join(ROOT, "docs", "CSI_Progress_Update.docx")

BLACK = RGBColor(0, 0, 0)
FONT = "Times New Roman"


def _style(run, size=12, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = BLACK


def para(doc, text="", size=12, bold=False, italic=False, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if text:
        _style(p.add_run(text), size=size, bold=bold, italic=italic)
    return p


def heading(doc, text, size=13):
    p = para(doc, text, size=size, bold=True, space_after=4)
    p.paragraph_format.space_before = Pt(10)
    return p


def bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_lead:
        _style(p.add_run(bold_lead), bold=True)
        _style(p.add_run(text))
    else:
        _style(p.add_run(text))
    return p


def figure(doc, filename, caption, width=5.8):
    path = os.path.join(ASSETS, filename)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=Inches(width))
    cap = para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    _style(cap.add_run(caption), size=10, italic=True)


doc = Document()

# Base style: Times New Roman 12 black everywhere.
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(12)
normal.font.color.rgb = BLACK

# --- Title block ---
t = para(doc, "Wi-Fi CSI Presence & Activity Sensing", size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para(doc, "Progress Update", size=13, bold=True,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para(doc, "July 13, 2026", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

# --- Summary ---
heading(doc, "Summary")
para(doc, "Over the past two weeks we extended the single-link CSI study from one "
     "room to three, pooling all strong-link recordings into a single model and "
     "evaluating it under leave-one-placement-out cross-validation. The single "
     "transmitter-to-receiver link is now fully characterized: we have mapped what "
     "it can do, what it cannot, and the hardware limit that governs both. The "
     "headline results are that occupancy detection is portable across placements, "
     "fine activity recognition is strong on-site but does not yet transfer to a new "
     "placement without a short calibration, and the usable range of a single link "
     "is a hard limit at roughly 120 inches. These findings define the deployment "
     "recipe and motivate the next phase (a second and third receiver).")

# --- Dataset ---
heading(doc, "Dataset")
para(doc, "All results below use one subject across Library Study Room 320X, "
     "spanning multiple node placements and orientations.", space_after=4)
bullet(doc, "14 configurations (placement x orientation), 29 recording sessions.")
bullet(doc, "~2.7 hours of recording, 910,071 CSI packets logged.")
bullet(doc, "4,971 labeled 2-second analysis windows.")
bullet(doc, "Classes: empty, stand, sit, walk, run.")
para(doc, "Two weak-link configurations were excluded because the radio link had "
     "degraded past its usable range (see Figure 5).", space_after=4)

# --- Key discoveries ---
heading(doc, "Key Discoveries")
bullet(doc, "Occupancy detection is portable. Presence (empty vs. occupied) holds "
       "up when the model is dropped into a placement it never trained on "
       "(0.79 balanced accuracy held-out, ~0.97 when calibrated on-site).",
       bold_lead="")
bullet(doc, "Fine activity needs on-site calibration. Activity classification is "
       "excellent when calibrated at the location it runs (0.78 balanced) but falls "
       "to 0.42 across unseen placements. A short empty-room baseline recovers it.")
bullet(doc, "Link quality, not the room, drives generalization. Strong links "
       "(RSSI >= ~-52 dBm) generalize across placements; weak links do not. This "
       "explained an earlier failure in a large open room.")
bullet(doc, "A single link has a hard range limit of ~120 inches (~10 ft). Beyond "
       "that, signal strength falls below ~-60 dBm and packets drop heavily.")
bullet(doc, "The long-block recording protocol fixed the chronically weak classes. "
       "Recording each activity as one continuous 2-minute block gives balanced, "
       "clean data and raises stand, sit, and run recall substantially on-site.")

# --- Figures ---
heading(doc, "Results")

figure(doc, "fig1_confusion_5class.png",
       "Figure 1. Activity confusion matrix (cross-placement, held-out). Rows are "
       "the true activity and sum to 100%. Empty is reliably identified (83%); "
       "stand/sit confuse with each other and run is largely absorbed into walk.")

figure(doc, "fig2_recall_5class.png",
       "Figure 2. Per-activity recall (cross-placement). Empty (0.83) and walk "
       "(0.55) transfer well; stand (0.31), sit (0.30) and run (0.10) are the weak "
       "classes when generalizing to a new placement.")

figure(doc, "fig3_perinstall_vs_crossplacement.png",
       "Figure 3. Per-install vs. cross-placement. Presence stays strong either way "
       "(0.97 -> 0.79), so occupancy is portable. Activity is excellent on-site "
       "(0.78) but drops to 0.42 across unseen placements.")

figure(doc, "fig4_confusion_presence.png",
       "Figure 4. Presence confusion matrix (empty vs. occupied). Occupied rooms "
       "are caught almost always; the residual error is occasionally calling a "
       "truly empty room occupied, driven by a motionless person.")

figure(doc, "fig5_link_range.png",
       "Figure 5. Link range limit. Below ~120 inches the link is strong and the "
       "model generalizes (green); beyond ~120-125 inches RSSI drops below ~-60 dBm, "
       "packets drop, and performance collapses (red).")

figure(doc, "fig6_confusion_blocks.png",
       "Figure 6. Activity confusion, long-block protocol, per-install (calibrated "
       "on-site). The diagonal is strong across the board (stand 92%, sit 90%, run "
       "73%, empty 81%); the only notable confusion is walk vs. run. Balanced "
       "accuracy 0.81.")

# --- Interpretation ---
heading(doc, "What This Means")
para(doc, "The single-link characterization is essentially complete. One antenna "
     "sees the room from a single geometry, which is enough to answer \"is someone "
     "present?\" robustly and across placements, and enough to classify activities "
     "well when calibrated on-site. The same single viewpoint is what limits "
     "cross-placement activity transfer, causes the stand-vs-sit and run-vs-walk "
     "confusions, and makes a motionless person hard to separate from an empty room.")

# --- Deployment recipe ---
heading(doc, "Deployment Guidance")
bullet(doc, "Keep the two nodes within ~120 inches and verify RSSI is at least "
       "~-55 dBm before recording.")
bullet(doc, "Record a short empty-room baseline (~90 seconds) at each install; it is "
       "required for reliable activity and recommended for presence.")
bullet(doc, "Use the long-block recording protocol for activity data collection.")

# --- Next steps ---
heading(doc, "Next Steps")
bullet(doc, "Add a second and third receiver. Independent viewpoints directly "
       "target coverage, cross-placement activity transfer, and the fine-activity "
       "confusions. The spare hardware is on hand.")
bullet(doc, "Validate the multi-receiver capture path (access point as receiver plus "
       "a promiscuous sniffer on the transmitter's unicast frames), then re-run this "
       "analysis with two and three viewpoints.")
bullet(doc, "Extend beyond one subject to add cross-subject generalization and "
       "formal confidence intervals.")

doc.save(OUT)
print("wrote", OUT)
