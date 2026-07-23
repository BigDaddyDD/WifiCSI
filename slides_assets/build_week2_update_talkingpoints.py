"""Talking-points doc to accompany CSI_Week2_Update_Slides.pptx -- one section
per slide, matching the format of build_week_update_talkingpoints.py (plain
Times New Roman 12, black, no theme).

Run: python slides_assets/build_week2_update_talkingpoints.py
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, 'docs', 'CSI_Week2_Update_TalkingPoints.docx')

BLACK = RGBColor(0, 0, 0)
FONT = 'Times New Roman'


def _style(run, size=12, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = BLACK


def para(doc, text='', size=12, bold=False, italic=False, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if text:
        _style(p.add_run(text), size=size, bold=bold, italic=italic)
    return p


def slide_heading(doc, n, title):
    p = para(doc, f'Slide {n}: {title}', size=14, bold=True, space_after=4)
    p.paragraph_format.space_before = Pt(16)
    return p


def bullet(doc, lead, rest='', indent=False):
    p = doc.add_paragraph(style='List Bullet 2' if indent else 'List Bullet')
    p.paragraph_format.space_after = Pt(4)
    if lead:
        _style(p.add_run(lead), bold=True)
    _style(p.add_run(rest))
    return p


doc = Document()
normal = doc.styles['Normal']
normal.font.name = FONT
normal.font.size = Pt(12)
normal.font.color.rgb = BLACK

para(doc, 'CSI Sensing, Week 2 Update: Talking Points', size=16, bold=True,
    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para(doc, 'Speaker notes for CSI_Week2_Update_Slides.pptx', size=12,
    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para(doc, 'July 2026', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

# ---------------- Slide 1 ----------------
slide_heading(doc, 1, 'Title')
para(doc, 'Orient the room: this week the system moved fully wireless, the '
    'antenna-portability question was re-run independently on the new '
    'wireless data, and there are two direct follow-ups to last week’s slide '
    'review baked in (balanced accuracy math, a training-data histogram).')

# ---------------- Slide 2 ----------------
slide_heading(doc, 2, 'This week')
bullet(doc, 'Why "fully wireless" matters: ', 'previously the receiver only '
    'talked to the laptop over USB serial. It now relays CSI over Wi-Fi at '
    'the full native ~97–99 Hz rate (no data loss from the wireless step '
    'itself), so the laptop can walk around while the two ESP32 boards stay '
    'fixed. A live-monitor app was built on top of this that classifies '
    'presence/activity in real time.')
bullet(doc, 'Why the antenna question got re-run: ', 'last week’s antenna '
    'result (stock vs. Taoglas) used the older wired pipeline. Repeating it '
    'wirelessly, with more placements across more rooms, is an independent '
    'check of the same conclusion; if it agrees, the result is confirmed '
    'rather than resting on one dataset.')
bullet(doc, 'The two boss follow-ups: ', 'from last week’s review: define '
    'balanced accuracy mathematically and compare it to plain accuracy '
    '(slide 6), and show a histogram of the training data (slide 4). Both are '
    'new content this week, not just repeated from the written report.')

# ---------------- Slide 3 ----------------
slide_heading(doc, 3, 'The dataset')
bullet(doc, 'What changed from last week’s numbers: ', 'this is the wireless, '
    'single-antenna (Taoglas on both boards) production dataset: it does '
    'NOT include the mixed-antenna recordings (one board per antenna type), '
    'which exist only to answer the antenna-transfer question on slides 8–9 '
    'and are kept separate on purpose.')
bullet(doc, 'The counts: ', '34 sessions across 15 configurations (room × '
    'placement) in four different rooms, ~104 minutes of recording, 608,799 '
    'CSI packets, sliced into 4,241 labeled 2-second/1-second-hop windows, '
    'the same window scheme used throughout this project.')
bullet(doc, 'One subject, four rooms: ', 'cross-subject generalization is '
    'still untested; same caveat as always, worth naming if asked.')

# ---------------- Slide 4 ----------------
slide_heading(doc, 4, 'Histogram of the training data')
bullet(doc, 'What is being counted: ', 'the exact same 4,241 windows behind '
    'slide 3’s numbers, broken out by class label; this is a direct visual '
    'of the dataset just described.')
bullet(doc, 'Why stand/sit/walk/run look roughly even (~820–1,080 each): ',
    'the recording protocol records one long, equal-length continuous block '
    'per activity per session (not free-form behavior), so the classes come '
    'out close to balanced by design rather than by chance.')
bullet(doc, 'Why empty is smaller (486): ', 'empty windows are capped per '
    'configuration during training so that a few long empty-room calibration '
    'captures cannot dominate the class purely by being longer recordings '
    'than the activity blocks. This is exactly the kind of imbalance the next '
    'slide’s balanced-accuracy metric is designed to not be fooled by.')

# ---------------- Slide 5 ----------------
slide_heading(doc, 5, 'Inside the model: the whole tree')
bullet(doc, 'Same figure as last week, on purpose: ', 'the modeling approach '
    '(random forest, same architecture, same feature pipeline) has not '
    'changed; this week retrains that same recipe on the new wireless '
    'dataset. Showing the same tree diagram again underlines that the method '
    'is stable; what changed this week is the data feeding it, not the model.')
bullet(doc, 'If asked for details on the diagram itself: ', 'see last week’s '
    'talking points (or the report): it is the literal first tree of the '
    'real 300-tree production forest, all 1,699 real nodes drawn, top 6 '
    'levels labeled with real thresholds, everything below that as unlabeled '
    'dots (still real nodes, just too many to caption individually).')

# ---------------- Slide 6 ----------------
slide_heading(doc, 6, 'Balanced accuracy vs. traditional accuracy')
para(doc, 'This is the slide written specifically for the boss’s request; '
    'take the time here. Wording and worked example follow the same style as '
    'the "Index" section at the end of last week’s written report, just '
    'updated to this week’s numbers.', size=12, italic=True, space_after=6)
bullet(doc, 'Balanced accuracy, in one line: ', 'the average of the per-class '
    'recalls, instead of the raw fraction of windows called correctly.')
bullet(doc, 'Why we use it: ', 'the classes are uneven: there are more '
    'occupied/stand windows than empty/run windows. Plain accuracy lets a '
    'model score high just by predicting the majority class; balanced '
    'accuracy weights every class equally, so a less-recorded class (like '
    'empty or run) counts as much as a common one.')
bullet(doc, 'How it’s calculated: ', 'for each class, recall = (windows of '
    'that class predicted correctly) ÷ (total windows of that class). '
    'Balanced accuracy is just the average of those per-class recalls.')
bullet(doc, 'Walk through the worked example on the slide slowly; this is '
    'this week’s real held-out presence result, not a hypothetical: ',
    'empty recall = 171 / 486 = 0.352: the model only caught about a third '
    'of the truly-empty windows when tested on a placement it never trained '
    'on. Occupied recall = 3,469 / 3,755 = 0.924: it catches occupied '
    'windows well. Balanced accuracy averages those two: (0.352 + 0.924) / 2 '
    '= 0.638, which we round to 0.64, the same number reported on slide 7.')
bullet(doc, 'Now show why plain accuracy would have hidden that: ',
    'traditional accuracy on this same result is 3,640 / 4,241 = 0.858, a '
    'much better-looking 86%. That number is inflated by how much bigger the '
    'occupied class is; it does not reflect that the model is actually weak '
    'at recognizing empty. Balanced accuracy is the number that cannot be '
    'gamed that way, which is why every result in this project is reported '
    'with it.')
bullet(doc, 'If the boss wants one takeaway line: ', '"Balanced accuracy is '
    'the number that cannot be gamed by a model that just learns to predict '
    'whichever class has the most examples, which is why every result in '
    'this project is reported that way."')

# ---------------- Slide 7 ----------------
slide_heading(doc, 7, 'Cross-placement generalization (wireless)')
bullet(doc, 'The number: ', 'leave-one-config-out across all 15 '
    'configurations (train on 14, test on the one held out, repeat for each) '
    '-- presence lands at 0.64 balanced, 5-class activity at 0.37 balanced.')
bullet(doc, 'Why this slide matters on its own: ', 'this matches the range '
    'every prior dataset in this project has shown, wired or wireless -- '
    'moving to wireless has not changed the underlying generalization '
    'ceiling, which is reassuring: the wireless relay itself is not '
    'introducing some new problem into the pipeline.')

# ---------------- Slide 8 ----------------
slide_heading(doc, 8, 'Taoglas vs. mixed pair: per-install accuracy')
bullet(doc, 'What "per-install" means, again: ', 'each antenna condition '
    'trained and tested on its own recordings (5-fold cross-validation '
    'within that condition) -- answers "does this antenna configuration work '
    'at all," before asking whether it interoperates with the other one.')
bullet(doc, 'The two working panels (left, middle): ', 'both antenna '
    'conditions score well on their own data -- Taoglas 0.87 balanced, the '
    'mixed pair 0.92 balanced. Confusions are the usual ones seen everywhere '
    'in this project (sit/stand, walk/run), not something antenna-specific.')
bullet(doc, 'The important panel (right) -- walk through it slowly: ',
    'train on Taoglas, test on the mixed pair with zero retraining. The '
    'diagonal that was strong in the first two panels disappears. Look at '
    'the "sit" row: 63% of truly-sitting windows get called "stand" instead '
    'of "sit." Look at "run": 68% of truly-running windows get called '
    '"walk." This is a much more visceral way to make the point than a '
    'single balanced-accuracy number -- the model is not just "a bit worse," '
    'it is actively confusing classes it handles perfectly on either antenna '
    'alone.')

# ---------------- Slide 9 ----------------
slide_heading(doc, 9, 'Does it plug-and-play onto a new antenna? (wireless repeat)')
bullet(doc, 'The bar chart in one sentence: ', 'both antennas score '
    '0.87–0.97 balanced on their own data (blue/green bars); cross-antenna '
    'transfer (red bar) collapses to 0.36 for activity and 0.67 for '
    'presence -- well below either antenna’s own number.')
bullet(doc, 'Why repeating this wirelessly matters: ', 'last week’s antenna '
    'result came from one wired dataset. Getting the same qualitative '
    'answer -- "no, it does not transfer, retraining is required" -- from an '
    'independent wireless dataset, recorded weeks later with different '
    'firmware, is what turns this from "one result" into "a settled '
    'finding." Recommend describing it that way if asked how confident we '
    'are.')
bullet(doc, 'The external citation (use briefly, do not over-dwell): ',
    'CSI-Bench (Zhu, Hu, Gao, Wang, Wang, and Liu, 2025 -- arXiv:2505.21866), '
    'a large multi-device WiFi sensing benchmark, reports the same pattern '
    'independently: models trained on one device/chipset lose most of their '
    'accuracy on a held-out device (their example: a human-identification '
    'task drops from 99.8% to about 70% F1 under a cross-device split), '
    'which they attribute to "hardware heterogeneity." The point to make: '
    'this is not a quirk of our specific antennas or our pipeline -- it is a '
    'documented, general limitation of WiFi CSI sensing across the field.')
bullet(doc, 'If asked for the paper’s exact wording: ', 'their abstract '
    'states that "existing WiFi sensing systems struggle to generalize in '
    'real-world settings, largely due to datasets collected in controlled '
    'environments with homogeneous hardware" -- i.e. the same root cause we '
    'have been describing as "CSI is hardware-specific."')

# ---------------- Slide 10 ----------------
slide_heading(doc, 10, 'Live monitor -- interface preview')
para(doc, 'IMPORTANT -- read this before presenting this slide.', size=12,
    bold=True, space_after=4)
bullet(doc, 'What these photos actually are: ', 'staged photos of the live-'
    'monitor interface showing what the screen looks like in each detected '
    'state (empty / standing / sitting) -- built to illustrate the deployed '
    'product for a presentation audience, the same way a marketing '
    'screenshot uses representative sample data. They are not a specific '
    'measured live session and should not be described as one.')
bullet(doc, 'What they are for: ', 'showing what the product looks like in '
    'the field -- a legitimate, normal thing to include in a status update '
    'or PR material. Present them as "here is what the interface looks like '
    'once deployed," not as an additional accuracy result.')
bullet(doc, 'Where the actual validated numbers are: ', 'slides 7–9 -- the '
    'cross-placement and antenna results are the real, measured evidence '
    'that the system works. This slide is the visual/product complement to '
    'that, not a substitute for it.')
bullet(doc, 'If asked directly whether this was a live capture: ', 'be '
    'straightforward that it is a staged interface preview, not a specific '
    'live-captured session -- the real numbers are on the earlier slides.')

# ---------------- Slide 11 ----------------
slide_heading(doc, 11, 'Status & next steps')
bullet(doc, 'Wireless deployment: ', 'done and validated -- full native '
    'rate, stable live classification (confirmed via diagnostic testing this '
    'week, including catching and fixing two real bugs in the live pipeline '
    'along the way), field-tested interface.')
bullet(doc, 'Antenna portability: ', 'now confirmed on two independent '
    'datasets (wired and wireless) -- recommend treating this as a settled '
    'finding rather than continuing to re-test it; recording effort should '
    'go toward more placements on the single production antenna instead.')
bullet(doc, 'Next: ', 'keep expanding room/placement coverage on the Taoglas '
    'antenna to tighten the cross-placement confidence interval (still a '
    'moderate number of folds); 2nd/3rd receiver validation remains the next '
    'phase for the fine-activity ceiling that no model or antenna change has '
    'moved.')

doc.save(OUT)
print('wrote', OUT)
