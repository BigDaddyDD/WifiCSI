"""Talking-points doc to accompany CSI_Week_Update_Slides.pptx — one section per
slide, so the presenter has the full explanation without it living on the slide
itself. Plain Times New Roman 12, black, no theme (matches the other docs).

Run: python slides_assets/build_week_update_talkingpoints.py
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, 'docs', 'CSI_Week_Update_TalkingPoints.docx')

BLACK = RGBColor(0, 0, 0)
FONT = 'Times New Roman'


def _style(run, size=12, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = BLACK


def para(doc, text='', size=12, bold=False, italic=False, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if text:
        _style(p.add_run(text), size=size, bold=bold, italic=italic)
    return p


def slide_heading(doc, n, title):
    p = para(doc, f'Slide {n} — {title}', size=14, bold=True, space_after=4)
    p.paragraph_format.space_before = Pt(16)
    return p


def bullet(doc, lead, rest='', indent=False):
    p = doc.add_paragraph(style='List Bullet 2' if indent else 'List Bullet')
    p.paragraph_format.space_after = Pt(4)
    if lead:
        _style(p.add_run(lead), bold=True)
    _style(p.add_run(rest))
    return p


doc = Document()
normal = doc.styles['Normal']
normal.font.name = FONT
normal.font.size = Pt(12)
normal.font.color.rgb = BLACK

para(doc, 'CSI Sensing — Week Update: Talking Points', size=16, bold=True,
    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para(doc, 'Speaker notes for CSI_Week_Update_Slides.pptx', size=12,
    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para(doc, 'July 2026', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

# ---------------- Slide 1 ----------------
slide_heading(doc, 1, 'Title')
para(doc, 'Nothing to present — just orient the room: this is a weekly check-in, '
    'not a final report. Two threads this week: (1) is random forest still the '
    'right model, and (2) does the model survive a hardware swap (different '
    'antenna, eventually a different capture device).')

# ---------------- Slide 2 ----------------
slide_heading(doc, 2, 'This week')
bullet(doc, 'Why this question came up: ', 'leadership asked whether the model '
    'would directly plug-and-play onto different hardware — a Taoglas antenna, '
    'or eventually an entirely different Wi-Fi adapter (e.g. an Alfa network '
    'card). That is a fair question because nothing in the pipeline had tested '
    'it yet — every prior result used the same stock antenna throughout.')
bullet(doc, 'What "same protocol" means: ', 'both experiments this week reuse '
    'the exact recording protocol from the main study (empty baseline + '
    'stand/sit/walk/run blocks) and the exact same feature pipeline and '
    'evaluation method as all previous results, so any difference we see is '
    'attributable to the thing being tested — not a change in methodology.')
bullet(doc, 'Two separate questions, don’t conflate them: ', 'the model '
    'comparison (slide 4) asks "is our algorithm choice good?" The antenna test '
    '(slides 7–8) asks "does a trained model survive a hardware change?" '
    'They use overlapping but not identical datasets.')

# ---------------- Slide 3 ----------------
slide_heading(doc, 3, 'The dataset')
bullet(doc, 'What these numbers are: ', 'the recordings that actually feed the '
    'main model (figures 7, 10, 11 and the earlier presence/activity results) — '
    'two library study rooms, 14 placement/orientation configurations, one '
    'subject. A separate weak-link room and one range-limited configuration were '
    'recorded but excluded (documented in the main report) because the radio '
    'link had degraded past a usable range, not because the numbers were '
    'inconvenient.')
bullet(doc, 'Recording time vs. packets vs. windows — three different counts, '
    'don’t conflate them: ', '"recording time" and "packets logged" describe '
    'the raw capture (~113 minutes, 669,353 CSI packets at ~97 Hz). "Labeled '
    'analysis windows" (4,971) is the number of 2-second, 1-second-hop slices '
    'the model actually trains and is scored on — many overlapping windows come '
    'from each recording, so this number is larger than "sessions" but smaller '
    'than "raw packets."')
bullet(doc, 'One subject is a stated limitation, not a hidden one: ', 'every '
    'number in this deck comes from a single person. Cross-subject '
    'generalization has not been tested yet and should be named as future work '
    'if asked — different body size/gait could plausibly change the fine-'
    'activity numbers more than the presence numbers.')
bullet(doc, 'Correction from an earlier count: ', 'a bug in the totals script '
    'was comparing the wrong folder level and silently swept in an unrelated '
    'room’s recordings; the dataset totals here are the corrected, verified '
    'numbers for the actual set of rooms the model trains on.')

# ---------------- Slide 4 ----------------
slide_heading(doc, 4, 'Why we still use a random forest')
bullet(doc, 'What was actually tested: ', 'ten model families — eight '
    'classical models on the same 160 hand-crafted features the pipeline '
    'already uses, plus a 1-D CNN and a GRU that instead learn directly from '
    'the raw calibrated CSI window. Everything else — data, calibration, '
    'train/test split — was held fixed, so the model is the only variable.')
bullet(doc, 'What the "160 engineered features" actually are: ', 'each '
    '2-second window is reduced to 160 numbers before any model sees it. '
    'For each of the ~52 usable subcarriers we compute three things: how far '
    'that subcarrier’s amplitude in this window differs from the empty-room '
    'baseline (its "relative amplitude change"), how much it fluctuates within '
    'the window (its "temporal standard deviation" — a motion proxy), and what '
    'fraction of its signal energy falls in the 0.5–5 Hz band typical of human '
    'movement (its "motion-band fraction," from an FFT of the window). That is '
    '3 × 52 = 156 numbers. Four more summarize the whole window at once (the '
    'mean and max of the relative-amplitude-change values, the overall '
    'deviation-from-empty distance, and the mean temporal std) for 156 + 4 = '
    '160 total. These are hand-designed, not learned — the alternative tested '
    'here is the two deep models, which skip this step and learn their own '
    'features straight from the raw amplitude instead.')
para(doc, 'One line on each model tested — what it is, and why it was worth '
    'including:', size=12, italic=True, space_after=4)
bullet(doc, 'RandomForest (our baseline): ', 'many decision trees, each trained '
    'on a random slice of data and features, averaged by majority vote. Robust '
    'to noisy/unscaled features and hard to overfit badly — the pipeline’s '
    'default since the start of the project.', indent=True)
bullet(doc, 'ExtraTrees: ', 'a more randomized cousin of random forest (it also '
    'randomizes the split thresholds, not just which features to consider). '
    'Tested to see whether more randomness reduces error further.', indent=True)
bullet(doc, 'GradientBoosting: ', 'builds trees one at a time, each new tree '
    'correcting the previous trees’ mistakes. Often the strongest tree-based '
    'method on tabular data — tested as the "best-in-class booster" comparison.',
    indent=True)
bullet(doc, 'LogisticRegression: ', 'the simplest model tried — a linear '
    'decision boundary per class. Fast and fully interpretable; included as the '
    'floor, to see how much the fancier models are actually buying us.', indent=True)
bullet(doc, 'SVM (RBF kernel): ', 'finds the boundary with the widest margin '
    'between classes, bent into nonlinear shapes via a kernel trick. A '
    'historically strong performer on small/medium tabular datasets like ours.',
    indent=True)
bullet(doc, 'k-NN (k=15): ', 'does not really "train" — it memorizes all '
    'examples and classifies a new window by vote of its 15 nearest labeled '
    'neighbors. Included as a sanity check that the engineered features '
    'meaningfully cluster by class at all.', indent=True)
bullet(doc, 'NaiveBayes: ', 'fits a simple probability distribution per class '
    'assuming features are independent. Practically free to train; the '
    'cheapest possible statistical baseline.', indent=True)
bullet(doc, 'MLP (small neural net): ', 'a basic feed-forward network on the '
    'same 160 features. Tests whether a neural net helps even without touching '
    'raw signal / deep learning.', indent=True)
bullet(doc, '1D-CNN (raw): ', 'a convolutional network fed the RAW calibrated '
    'amplitude window instead of engineered features — it slides filters across '
    'the time axis, so it learns its own "shape" detectors directly from the '
    'amplitude waveform. This directly tests the raw-signal deep-learning idea.',
    indent=True)
bullet(doc, 'GRU (raw): ', 'a recurrent network that reads the window as a time '
    'sequence, carrying a memory state as it scans across time. A second raw/'
    'deep-learning architecture, since RNNs are typically strong on '
    'sequential data.', indent=True)
bullet(doc, 'Why the other classical models fall short: ', 'they don’t, '
    'really — that is the point. Every classical model lands within a few '
    'points of the random forest on cross-placement activity (0.37–0.46 '
    'balanced accuracy) and several are competitive on presence. The random '
    'forest is not uniquely good; it is simply tied for best while also being '
    'fast to train, robust to unscaled features, and easy to interpret (hence '
    'the tree diagram on the next slide). Logistic regression and the RBF-SVM '
    'actually edge it out once the model is calibrated on-site (0.85 and 0.83 '
    'vs. 0.79 per-install) — worth mentioning if asked "could we simplify '
    'further?" Yes, for the deployed/calibrated case a linear model would work '
    'almost as well.')
bullet(doc, 'Why deep learning falls short here: ', 'the 1D-CNN and GRU learn '
    'their own features from raw amplitude instead of using the hand-crafted '
    'ones. The CNN is nominally the best cross-placement activity score '
    '(0.46) because it picks up walk/run motion dynamics well, but it is '
    'noticeably worse at presence (0.59 vs. 0.75–0.84 for the feature-'
    'based models) because it doesn’t "know" to compare against an empty '
    'baseline the way the engineered features are built to. Deep nets also need '
    'far more data than the ~5,000 windows we have to reliably beat a simpler '
    'model — this is a data-scale limitation, not a fundamental one. If '
    'the multi-receiver phase produces an order of magnitude more data, deep '
    'learning is worth revisiting.')
bullet(doc, 'The headline message: ', 'model choice barely moves the needle on '
    'the hard problem (cross-placement fine activity). That result argues the '
    'ceiling is the single-antenna sensor itself, not the algorithm sitting on '
    'top of it — which is the whole justification for the next hardware '
    'phase (2nd/3rd receiver) instead of continuing to chase a better model.')

# ---------------- Slide 5 ----------------
slide_heading(doc, 5, 'Inside the model: the whole tree')
bullet(doc, 'What this actually is: ', 'the literal first tree inside the real, '
    'deployed 300-tree forest (`clf.estimators_[0]`) — not a stand-in built for '
    'the slide. Every one of its 1,699 real nodes is drawn; nothing is cut off '
    'or approximated. It is genuinely deep: 36 levels, 850 leaves.')
bullet(doc, 'Why most of it has no text: ', 'a tree this size cannot have 1,699 '
    'readable text boxes in one image — there is no way around that and still '
    'see the whole shape at once. The top 6 levels (63 nodes) show their real '
    'split condition, the share of training data reaching that point, and the '
    'majority class; everything below that (1,636 nodes, down to level 36) is '
    'still the exact same real tree, just drawn as a small colored dot instead '
    'of a labeled box, so the whole structure is visible without the image '
    'becoming unreadable clutter. Every line and every dot is a real split/leaf '
    '— none of it is simplified or invented, only unlabeled.')
bullet(doc, 'How the layout avoids overlapping text: ', 'each depth level is '
    'independently spaced across the same width, rather than spacing every '
    'node in proportion to how many leaves sit below it. A depth-1 split with '
    '400 descendants and a depth-1 split with 50 both get equal room at that '
    'row — this is what makes a 1,699-node tree fit in one image at all, since '
    'the real subtree sizes are wildly uneven.')
bullet(doc, 'The story the top of the tree tells: ', 'the very first question is '
    'essentially "how different does this window look from the empty room?" — '
    'that alone splits off a large empty-leaning branch. The next levels ask '
    'about motion-band energy and individual subcarrier deviations to start '
    'separating stationary (stand/sit) from moving (walk/run); the enormous '
    'fan below is the tree continuing to refine that distinction for another '
    '30 levels using finer and finer subcarrier detail.')
bullet(doc, 'Why 300 of these instead of one: ', 'a single tree this deep is '
    'prone to memorizing quirks of the specific training placements — 36 '
    'levels is enough to carve out a rule for almost every individual window '
    'it was trained on. Averaging 300 independently-grown trees (each on a '
    'different bootstrap sample and random feature subset) cancels out each '
    'tree’s individual overfitting, which is why the forest generalizes far '
    'better than this one tree does by itself.')

# ---------------- Slide 6 ----------------
slide_heading(doc, 6, 'Which features the model actually uses')
bullet(doc, 'What "feature importance" measures: ', 'for each of the 160 '
    'engineered numbers (explained on slide 4), this is how much that number '
    'contributes, on average across all 300 trees, to correctly splitting the '
    'classes. Bigger bar = the model leans on it more.')
bullet(doc, 'The top three are all "how different from empty" measures: ', '‘max '
    'relative amplitude change’, ‘mean relative amplitude change’, and '
    'the overall deviation-from-empty distance are the three most-used '
    'features, by a clear margin over any single subcarrier. This directly '
    'confirms the empty-baseline calibration recipe is doing most of the real '
    'work — not any particular exotic feature.')
bullet(doc, 'No single subcarrier dominates: ', 'after the top three summary '
    'features, importance is spread thinly across many individual '
    'subcarriers (around 1% each) rather than concentrated on one or two. That '
    'is the practical argument for using all ~52 usable subcarriers rather than '
    'a cheaper single-frequency measurement — the motion signal is spread '
    'across the channel bandwidth, not localized to one frequency.')

# ---------------- Slide 7 ----------------
slide_heading(doc, 7, 'Stock vs. Taoglas: per-install accuracy')
bullet(doc, 'What "per-install" means here: ', 'the model is trained and tested '
    'on the same antenna’s own recordings (5-fold cross-validation within '
    'that antenna’s data) — this answers "does this antenna work at '
    'all," before we ask whether it interoperates with the other antenna.')
bullet(doc, 'Both antennas work: ', 'the ESP32 stock antenna reaches 0.72 '
    'balanced accuracy, the Taoglas 0.77 — both well above guessing. '
    'Confusions look the same for both — sit/stand get mixed up (both '
    'stationary) and run bleeds into walk (both whole-body motion) — which '
    'is the same pattern seen everywhere else in the project, so this isn’t '
    'an antenna-specific weakness.')
bullet(doc, 'Taoglas is measurably the stronger antenna: ', 'its link ran about '
    '10 dB stronger RSSI than the ESP32 stock antenna in the same positions, '
    'and the slightly higher accuracy (0.77 vs. 0.72) tracks that. This is '
    'consistent with the project’s established link-quality finding '
    '(stronger link → better generalization) — nothing new here about '
    'physics, just a direct confirmation with a second antenna.')
bullet(doc, 'One data-quality caveat if asked: ', 'one stock-antenna recording '
    '(a placement’s activity blocks) failed outright — zero packets '
    'logged, a known serial-stall failure mode — so the stock antenna’s '
    'result rests on two placements instead of three. This is exactly the kind '
    'of gap the new live link-health monitor in the recording GUI is meant to '
    'catch before it happens again.')

# ---------------- Slide 8 ----------------
slide_heading(doc, 8, 'Does it plug-and-play onto a new antenna?')
bullet(doc, 'This is the direct answer to leadership’s question: ', 'no. '
    'Train the model on the ESP32 stock antenna’s data and test it — '
    'unmodified — on the Taoglas antenna’s data (and vice versa), and '
    '5-class activity collapses to 0.18 balanced accuracy and presence to '
    '0.49 — both indistinguishable from guessing. The model is not '
    'transferring at all.')
bullet(doc, 'Why this happens (the physical reason): ', 'CSI amplitude is a '
    'property of the exact hardware and antenna gain pattern, not just the '
    'room. The ESP32 stock antenna and the Taoglas antenna produce two '
    'different raw signal signatures for the identical physical scene, so a '
    'model trained to recognize the stock antenna’s "empty room" signature '
    'has never seen the Taoglas antenna’s version of "empty."')
bullet(doc, 'Recalibration vs. retraining — the distinction that matters here: ',
    '"calibration" just re-centers features against that config’s own empty-'
    'room capture (a normalization step) — both the working per-install result '
    'and the failing cross-antenna test DO this, always. "Retraining" means '
    'actually fitting new decision boundaries on labeled activity examples '
    '(stand/sit/walk/run) recorded with that antenna. The good Taoglas '
    'per-install number (0.77) came from retraining — 5-fold cross-validation '
    'on Taoglas’s OWN labeled recordings, so the classifier genuinely saw '
    'Taoglas activity examples during training. What collapses to chance is '
    'the scenario where the classifier is only recalibrated (new empty '
    'baseline) but never shown a single Taoglas activity example — i.e. reusing '
    'the stock-trained decision boundaries. So recalibrating alone is not '
    'enough; a new antenna needs its own labeled training data, not just its '
    'own empty capture. That is the precise, less convenient version of "needs '
    'retraining" worth saying plainly if asked.')
bullet(doc, 'What DOES transfer: ', 'recalibrating AND retraining on the SAME '
    'antenna at a new placement works fine for both — the stock antenna '
    'reaches 0.87 presence / 0.72 activity per-install, the Taoglas 0.78 '
    'presence / 0.77 activity, matching results seen throughout the main '
    'study. The antenna itself is not the problem; swapping it without '
    'collecting new labeled data on it is.')
bullet(doc, 'The takeaway for leadership: ', 'a new antenna, or eventually a new '
    'capture device, is not a drop-in replacement for the current model. What '
    'transfers to new hardware is the collection protocol, the feature '
    'pipeline, and the labeling scheme — not the trained weights. Any new '
    'device needs its own labeled training data (for full accuracy); a '
    'calibration pass alone is not sufficient. This should be set as an '
    'expectation before anyone assumes otherwise.')

# ---------------- Slide 9 ----------------
slide_heading(doc, 9, 'Status & next steps')
bullet(doc, 'Model choice question: ', 'closed for now. Random forest stays the '
    'default — it is tied for best, cheap to train, and interpretable. '
    'Revisit deep learning only once the dataset is an order of magnitude '
    'bigger (post multi-receiver).')
bullet(doc, 'Antenna portability question: ', 'answered directly and '
    'negatively for zero-shot transfer (recalibration alone does not work), '
    'positively for retraining on the new antenna’s own data. Recommend '
    'continuing to record a few more stock/Taoglas position pairs to firm up '
    'the confidence interval on this result before calling it final.')
bullet(doc, 'In progress: ', 'a live link-quality monitor was added to the '
    'recording GUI this week (packet rate / loss / RSSI banner, real-time) so a '
    'failed recording like the one behind slide 7’s caveat gets caught '
    'immediately instead of discovered during analysis.')
bullet(doc, 'Next phase: ', 'validate a 2nd receiver (spare ESP32-C3 hardware '
    'already on hand) — this is the lever that plausibly fixes both '
    'coverage dead zones and the fine-activity ceiling that no model or '
    'antenna swap has moved.')

para(doc, '', space_after=2)
para(doc, 'Reproducibility: run "python gen_antenna_figs.py" and '
    '"python gen_tree_figs.py" to regenerate figures 8–11; '
    '"python model_comparison.py figure" regenerates figure 7 from the saved '
    'results without retraining all 10 models.',
    size=10, italic=True)

doc.save(OUT)
print('wrote', OUT)
