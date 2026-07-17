"""Build a plain model-comparison report (boss request: "how would various models
perform — CNNs, deep learning, etc."). Times New Roman 12, black, no theme.

Reads docs/model_report_assets/model_comparison.json (produced by
model_comparison.py) so every number is the real, reproducible output.
Run: python slides_assets/build_model_comparison_doc.py
"""
import json
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ASSETS = os.path.join(ROOT, "docs", "model_report_assets")
OUT = os.path.join(ROOT, "docs", "CSI_Model_Comparison.docx")
RESULTS = json.load(open(os.path.join(ASSETS, "model_comparison.json")))

BLACK = RGBColor(0, 0, 0)
FONT = "Times New Roman"

# input representation per model
FEAT = "160-D features"
RAW = "Raw window"
INPUT = {"1D-CNN (raw)": RAW, "GRU (raw)": RAW}


def _style(run, size=12, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = BLACK


def para(doc, text="", size=12, bold=False, italic=False, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if text:
        _style(p.add_run(text), size=size, bold=bold, italic=italic)
    return p


def heading(doc, text, size=13):
    p = para(doc, text, size=size, bold=True, space_after=4)
    p.paragraph_format.space_before = Pt(10)
    return p


def bullet(doc, lead, rest=""):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if lead:
        _style(p.add_run(lead), bold=True)
    _style(p.add_run(rest))
    return p


def figure(doc, filename, caption, width=6.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(os.path.join(ASSETS, filename), width=Inches(width))
    cap = para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    _style(cap.add_run(caption), size=10, italic=True)


def fmt(v):
    return "—" if v is None or v != v else f"{v:.2f}"     # em dash for NaN


doc = Document()
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(12)
normal.font.color.rgb = BLACK

para(doc, "Wi-Fi CSI Sensing — Model Comparison", size=16, bold=True,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para(doc, "How different model families perform on the same task", size=12,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para(doc, "July 13, 2026", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

# --- Purpose / method ---
heading(doc, "What was tested and how")
para(doc, "We compared ten model families on the occupancy/activity task, from "
     "simple classifiers to deep neural networks (a 1-D convolutional network and "
     "a recurrent GRU network). To make the comparison fair, everything except the "
     "model was held fixed: the same recordings, the same per-configuration "
     "empty-baseline calibration, and the same leave-one-configuration-out "
     "evaluation (train on all placements but one, test on the held-out placement). "
     "Any difference in the numbers is therefore due to the model itself, not the "
     "setup. Dataset: 14 strong-link configurations, 4,971 labeled 2-second windows.")
para(doc, "Two input representations were compared. The eight classical models "
     "receive the same 160 hand-crafted features per window used by the current "
     "pipeline. The two deep networks instead receive the raw calibrated window "
     "(52 subcarriers x 200 samples) and learn their own features end-to-end — "
     "this is the fairest test of whether deep learning helps here.")

# --- Table ---
heading(doc, "Results (balanced accuracy, held-out unless noted)")

order = sorted(RESULTS.keys(), key=lambda n: RESULTS[n]["act5_bal"], reverse=True)
tbl = doc.add_table(rows=1, cols=5)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = "Table Grid"
hdr = ["Model", "Input", "Presence\n(cross-placement)",
       "Activity 5-class\n(cross-placement)", "Activity 5-class\n(per-install)"]
for j, h in enumerate(hdr):
    c = tbl.rows[0].cells[j]
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _style(c.paragraphs[0].add_run(h), size=10, bold=True)
for n in order:
    r = RESULTS[n]
    row = tbl.add_row().cells
    vals = [n, INPUT.get(n, FEAT), fmt(r["presence_bal"]), fmt(r["act5_bal"]),
            fmt(r.get("per_install_bal"))]
    for j, v in enumerate(vals):
        cell = row[j]
        cell.paragraphs[0].alignment = (WD_ALIGN_PARAGRAPH.LEFT if j == 0
                                        else WD_ALIGN_PARAGRAPH.CENTER)
        _style(cell.paragraphs[0].add_run(v), size=10,
               bold=(n == "RandomForest (baseline)"))
para(doc, "", space_after=2)
para(doc, "RandomForest is the current pipeline's model (bold). “Cross-placement” "
     "is the honest generalization number (unseen placement); “per-install” is the "
     "deployable ceiling when the model is calibrated at the location it runs. Deep "
     "networks were not run per-install (they need far more data to be worth it). "
     "Chance for 5-class is 0.20; for presence 0.50.", size=10, italic=True)

figure(doc, "fig7_model_comparison.png",
       "Figure 7. Cross-placement balanced accuracy for every model, sorted by "
       "activity score. Presence (blue) and 5-class activity (orange).")

# --- Findings ---
heading(doc, "What the comparison shows")
bullet(doc, "Model choice barely changes cross-placement activity. ",
       "Every model lands between 0.37 and 0.46 balanced accuracy on an unseen "
       "placement — the 1-D CNN is nominally best (0.46), naive Bayes and gradient "
       "boosting next (0.45, 0.44), the random forest 0.42 — but the spread is "
       "within noise. No algorithm breaks out. The ceiling here is set by the "
       "single-antenna viewpoint and the amount of data, not by the model. A fancier "
       "model does not fix cross-placement activity; a second receiver is the lever.")
bullet(doc, "For presence, engineered features beat raw deep learning decisively. ",
       "The calibrated hand-crafted features give 0.79–0.84 presence with essentially "
       "any classifier, while the raw-window CNN and GRU reach only 0.59 and 0.64. "
       "The empty-baseline calibration already encodes “is the room different from "
       "empty?”; a deep net trained from scratch on this much data does not "
       "rediscover it as well.")
bullet(doc, "The deep networks captured motion but missed stillness. ",
       "The CNN had the best walk and run recall of any model (0.67 and 0.54) yet the "
       "worst empty and sit recall — it learns movement dynamics from the raw signal "
       "but confuses an empty room with a still person. The feature-based models are "
       "the opposite: they nail empty. This is a useful pointer for combining them "
       "later.")
bullet(doc, "Once calibrated on-site, the simplest models win. ",
       "Per-install, logistic regression (0.85) and the RBF-SVM (0.83) actually beat "
       "the tree ensembles (0.79–0.80). With an on-site empty baseline the features "
       "become cleanly separable, so a light linear model is enough — no heavy model "
       "needed for the realistic deployment mode.")
bullet(doc, "Deep learning does not beat the random forest at this data scale. ",
       "With ~5,000 windows from one subject, the CNN and GRU do not outperform the "
       "random forest overall (they trade a hair of activity accuracy for a large "
       "loss on presence) and cost far more compute. Deep models become worth "
       "revisiting once multi-receiver collection scales the dataset up by an order "
       "of magnitude.")

# --- Recommendation ---
heading(doc, "Recommendation")
para(doc, "Keep the random forest as the default model: it is at or near the top for "
     "presence (0.84), competitive on activity, needs no feature scaling or tuning, "
     "trains in seconds, and is interpretable. The result to take away is that the "
     "current limitations are a sensing limitation, not a modeling one — no model in "
     "this sweep, including deep networks, removes the cross-placement activity wall. "
     "That is the direct argument for adding a second and third receiver rather than "
     "investing in a heavier model on a single link.")

para(doc, "", space_after=2)
para(doc, "Reproducibility: run “python model_comparison.py” (classical + deep) to "
     "regenerate the table, Figure 7, and model_comparison.json. The random-forest "
     "row reproduces the main report's numbers exactly, which anchors the comparison.",
     size=10, italic=True)

doc.save(OUT)
print("wrote", OUT)
