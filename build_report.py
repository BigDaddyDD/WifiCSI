#!/usr/bin/env python3
"""Assemble the brief Phase-A interim report (.docx) from the generated
figures + metrics.json. Run report_figs.py first."""
import json
import os

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

A = os.path.join('docs', 'report_assets')
M = json.load(open(os.path.join(A, 'metrics.json')))
pc = M['presence']['calibrated']
pr = M['presence']['raw']
ac = M['activity']['calibrated']
rec = M['activity']['per_class_recall']
sens = M['baseline_sensitivity']
fold_lo = min(a for _, a in pc['folds']); fold_hi = max(a for _, a in pc['folds'])

doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)


def fig(name, caption, width=6.2):
    doc.add_picture(os.path.join(A, name), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


t = doc.add_heading('Wi-Fi CSI Presence & Activity Sensing', level=0)
sub = doc.add_paragraph()
sub.add_run('Phase A interim note — single-link characterization (July 2026)').italic = True

doc.add_heading('1. Setup', level=1)
doc.add_paragraph(
    "We are testing whether ordinary Wi-Fi hardware can tell when a room is "
    "occupied, and eventually what the person is doing, using Channel State "
    "Information (CSI) — the per-subcarrier channel measurement a receiver "
    "already computes for every packet. The current rig is a single link: two "
    "Seeed Studio XIAO ESP32-C3 boards (2.4 GHz, 20 MHz, single antenna). One "
    "board transmits a steady 100 Hz stream; the other receives it and reports "
    "CSI for 52 usable subcarriers at about 97 Hz to a laptop.")
doc.add_paragraph(
    "Data was collected in a small L-shaped library study room (roughly 12 ft "
    "across the short axis, open floor, furniture unchanged). For each of five "
    "configurations we ran the same scripted sequence: empty room, standing and "
    "sitting at marked floor spots, walking fixed back-and-forth paths, and "
    "running. The five configurations differ in where the two boards sit and how "
    "their antennas are oriented (flat / rotated 90° / vertical), because the "
    "system must not depend on one exact placement. That gives "
    f"{M['n_sessions']} sessions and {M['n_windows']:,} analysis windows.")

doc.add_heading('2. How we treated it as a data-science problem', level=1)
doc.add_paragraph(
    "Each recording is a stream of 52-subcarrier amplitude vectors at ~97 Hz. We "
    "cut it into 2-second windows (1-second overlap) and summarize each window "
    "with a fixed feature vector: how much each subcarrier fluctuates over the "
    "window, its average shape, and how much of its energy sits in the 0.5–5 Hz "
    "motion band. A random forest then classifies each window.")
doc.add_paragraph(
    "Two choices matter. First, raw CSI amplitude encodes the specific room "
    "geometry, so a model trained on it memorizes the setup instead of learning "
    "“person vs no person.” We therefore also compute calibrated features: each "
    "window expressed as its deviation from that configuration’s own empty-room "
    "baseline, then standardized. This is the deployment recipe — record a short "
    "empty baseline when you install the sensor, then measure change against it. "
    "Second, we evaluate with leave-one-session-out cross-validation: train on "
    "four configurations, test on the fifth one the model has never seen, and "
    "rotate. That is the honest measure of whether it works in a new setup. For "
    "contrast, a within-session random split (which lets the model peek at the "
    "test geometry) reaches about 0.83 on the 5-class task — most of that is "
    "memorization, which is exactly what the leave-one-session-out test removes.")

doc.add_heading('3. What it does well: presence detection', level=1)
doc.add_paragraph(
    f"With calibration, detecting whether the room is occupied generalizes to "
    f"configurations the model never saw: {pc['acc']*100:.0f}% overall accuracy "
    f"and {pc['bal']:.2f} balanced accuracy, with occupied correctly flagged "
    f"{M['presence']['occ_recall_calibrated']*100:.0f}% of the time and empty "
    f"{M['presence']['empty_recall_calibrated']*100:.0f}% of the time. Across the "
    f"five held-out configurations the accuracy ranged from {fold_lo:.2f} to "
    f"{fold_hi:.2f}, so no single placement is carrying the result.")
doc.add_paragraph(
    f"These numbers come straight from the leave-one-session-out predictions "
    f"pooled over all five held-out configurations. The comparison in Figure 1 "
    f"is the main point: without calibration the same features give only "
    f"{pr['bal']:.2f} balanced accuracy — the model labels almost everything "
    f"“occupied,” which looks acceptable on raw accuracy but is useless. "
    f"Calibration against the empty baseline is what turns CSI into a usable, "
    f"portable presence signal.")
fig('fig1_methods.png',
    'Figure 1. Leave-one-session-out performance. Left: presence. Right: 5-class '
    'activity. Blue = accuracy, red = balanced accuracy; dashed line = chance / '
    'majority-class. Calibration is what makes presence work.')
fig('fig2_presence_confusion.png',
    'Figure 2. Presence confusion matrix (calibrated, pooled over held-out '
    'configurations). Occupied is caught reliably; the harder case is confirming '
    'the room is truly empty.', width=4.3)

doc.add_heading('4. Where it falls short: fine activity classes', level=1)
doc.add_paragraph(
    f"Telling the four activities apart is a different story. On the 5-class task "
    f"{{empty, stand, sit, walk, run}} the calibrated model reaches "
    f"{ac['acc']*100:.0f}% accuracy ({ac['bal']:.2f} balanced) on unseen "
    f"configurations. It is still the best option — an RSSI-only baseline "
    f"(signal strength, no CSI) gets {M['activity']['rssi']['bal']:.2f} and raw "
    f"CSI {M['activity']['raw']['bal']:.2f} — but it is well short of reliable.")
doc.add_paragraph(
    f"The per-class breakdown (Figure 3) shows where it breaks. Empty "
    f"({rec['empty']:.2f}), standing ({rec['stand']:.2f}) and walking "
    f"({rec['walk']:.2f}) are recognizable, but sitting ({rec['sit']:.2f}) and "
    f"running ({rec['run']:.2f}) are not. Sitting gets confused with standing — "
    f"both are near-stationary and a single 2.4 GHz link sees little difference "
    f"between them — and running gets confused with walking, since both are "
    f"whole-body motion (running also has the least data so far, one segment per "
    f"configuration). These recalls are again from the pooled leave-one-session-out "
    f"predictions. The takeaway: one link is enough for presence but not for "
    f"clean activity separation, which is the main argument for adding a second "
    f"and third receiver to cover motion the single link misses.")
fig('fig3_perclass_recall.png',
    'Figure 3. Per-activity recall on unseen configurations. Green ≥ 0.5, red '
    'below. Sitting and running are the failure modes.', width=5.3)

doc.add_heading('5. A measurement finding worth flagging', level=1)
doc.add_paragraph(
    f"Because presence leans on the empty-room baseline, we checked whether that "
    f"baseline is stable. Within a session it is essentially fixed — the empty "
    f"fingerprint at the start and end of an eight-minute session correlates above "
    f"0.999, with signal strength moving under 1 dB. But how you capture the "
    f"baseline matters a lot. Calibrating from one short empty snapshot "
    f"underestimates the normal window-to-window fluctuation of an empty room and "
    f"produces false “occupied” readings (empty recall drops to "
    f"{sens['contiguous']['empty_recall']:.2f}, balanced accuracy "
    f"{sens['contiguous']['balanced']:.2f}). Sampling the empty baseline over a "
    f"longer spread of time recovers it "
    f"({sens['spread']['empty_recall']:.2f} / {sens['spread']['balanced']:.2f}). "
    f"The practical rule for deployment: record a longer, or periodically "
    f"refreshed, empty baseline at install rather than a single short capture.")
fig('fig4_baseline_sensitivity.png',
    'Figure 4. Same model and data, only the empty-calibration window changes. A '
    'short one-shot baseline underestimates normal variation and over-reports '
    'occupancy.', width=5.3)

doc.add_heading('6. Status and next steps', level=1)
doc.add_paragraph(
    "This is an interim snapshot, not the final study. So far it is one room, one "
    "subject, and one recording per configuration, so the numbers do not yet have "
    "confidence intervals — more sessions are being recorded and will tighten "
    "them. The clear reads today: presence detection is robust and transfers "
    "across placements once calibrated; fine activity classification is not there "
    "with a single link. Planned next steps are more configurations (for "
    "confidence intervals), more sitting and running data, then a second and "
    "third receiver to improve motion coverage, followed by other rooms and "
    "subjects.")

out = 'docs/Phase_A_Interim_Report.docx'
doc.save(out)
print('Saved', out)
